import openai
from docx import Document
from pydub import AudioSegment


def transcribe_audio(audio_file_path):
  with open(audio_file_path, 'rb') as audio_file:
    transcription = openai.Audio.transcribe(model='whisper-1', file=audio_file)
  return transcription#['text']


def meeting_minutes(transcription):
  abstract_summary = abstract_summary_extraction(transcription)
  key_points = key_points_extraction(transcription)
  action_items = action_item_extraction(transcription)
  sentiment = sentiment_analysis(transcription)

  return {
      'abstract_summary': abstract_summary,
      'key_points': key_points,
      'action_items': action_items,
      'sentiment': sentiment
  }


def abstract_summary_extraction(transcription):
    response = openai.ChatCompletion.create(
        model="gpt-4",
        temperature=0,
        messages=[
            {
                "role": "system",
                "content": """You are a highly skilled AI trained in language comprehension and summarization. \
                I would like you to read the following text and summarize it into a concise abstract paragraph. \
                Aim to retain the most important points, providing a coherent and readable summary \
                that could help a person understand the main points of the discussion without needing to read the entire text. \
                Please avoid unnecessary details or tangential points."""
            },
            {
                "role": "user",
                "content": transcription
            }
        ]
    )
    return response['choices'][0]['message']['content']


def key_points_extraction(transcription):
    response = openai.ChatCompletion.create(
        model="gpt-4",
        temperature=0,
        messages=[
            {
                "role": "system",
                "content": "You are a proficient AI with a specialty in distilling information into key points. \
                Based on the following text, identify and list the main points that were discussed or brought up. \
                These should be the most important ideas, findings, or topics that are crucial to the \
                essence of the discussion. Your goal is to provide a list that someone could read to quickly understand \
                what was talked about."
            },
            {
                "role": "user",
                "content": transcription
            }
        ]
    )
    return response['choices'][0]['message']['content']



def action_item_extraction(transcription):
    response = openai.ChatCompletion.create(
        model="gpt-4",
        temperature=0,
        messages=[
            {
                "role": "system",
                "content": "You are an AI expert in analyzing conversations and extracting action items. \
                Please review the text and identify any tasks, assignments, or actions that were agreed upon or \
                mentioned as needing to be done. These could be tasks assigned to specific individuals, or \
                general actions that the group has decided to take. Please list these action items clearly and concisely."
            },
            {
                "role": "user",
                "content": transcription
            }
        ]
    )
    return response['choices'][0]['message']['content']



def sentiment_analysis(transcription):
    response = openai.ChatCompletion.create(
        model="gpt-4",
        temperature=0,
        messages=[
            {
                "role": "system",
                "content": "As an AI with expertise in language and emotion analysis, your task is \
                to analyze the sentiment of the following text. Please consider the overall tone of the discussion, \
                the emotion conveyed by the language used, and the context in which words and phrases are used. \
                Indicate whether the sentiment is generally positive, negative, or neutral, and provide brief explanations \
                for your analysis where possible."
            },
            {
                "role": "user",
                "content": transcription
            }
        ]
    )
    return response['choices'][0]['message']['content']


def save_as_docx(minutes, filename):
  doc = Document()

  for key, value in minutes.items():
    # Replace underscores with spaces and capitalize each word for the heading
    heading = ' '.join(word.capitalize() for word in key.split('_'))
    doc.add_heading(heading, level=1)
    doc.add_paragraph(value)
    # Add a line break between sections
    doc.add_paragraph()

  doc.save(filename)


# input_key = input('Please enter your OpenAI API Key: ')
input_key = " "
openai.api_key = input_key




audio_file_path = 'sampleAudio.mp3'
# 2:53 mins long in wav format and 50.7 MB

audio_file_segment = AudioSegment.from_mp3(audio_file_path)

# PyDub handles time in milliseconds
ten_minutes = 3 * 60 * 1000

first_3_minutes = audio_file_segment[:ten_minutes]

first_3_minutes.export("sample.mp3", format="mp3")
# 5:p51 mins long in mp3 format and 2.8 MB


audio_file_path = 'sample.mp3'

transcription = transcribe_audio(audio_file_path)


print(transcription)

minutes = meeting_minutes(transcription['text'])

print(minutes)


save_as_docx(minutes, 'sampleAudio.docx')




